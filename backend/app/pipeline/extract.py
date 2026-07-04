"""Pipeline steps 3-4 — text extraction with page anchoring + OCR fallback.

- PDF: native text via PyMuPDF; pages that come back near-empty are treated as
  scanned images and re-read via Tesseract (``fra``).
- DOCX: python-docx (paragraphs + tables), single logical page.
- XLSX: openpyxl, one logical page per worksheet.

Robustness (CDC §6): a failed OCR on one page never sinks the rest of the DCE.
The page is kept with whatever text we have and flagged in ``warnings``.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

from app.config import Settings

log = logging.getLogger(__name__)


@dataclass
class PageText:
    page_number: int  # 1-indexed
    text: str
    ocr_used: bool = False


@dataclass
class ExtractResult:
    pages: list[PageText] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _ocr_page(page, settings: Settings) -> str:
    """Render a PDF page to an image and OCR it. Returns "" on any failure."""
    try:
        import pytesseract
        from PIL import Image

        pix = page.get_pixmap(dpi=settings.ocr_dpi)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img, lang=settings.ocr_lang) or ""
    except Exception:  # noqa: BLE001 — OCR is best-effort
        log.exception("OCR failed on page %s", getattr(page, "number", "?"))
        return ""


def extract_pdf(content: bytes, filename: str, settings: Settings) -> ExtractResult:
    import fitz  # PyMuPDF

    result = ExtractResult()
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"{filename}: PDF illisible ({exc})")
        return result

    with doc:
        for index in range(doc.page_count):
            page = doc.load_page(index)
            page_no = index + 1
            native = (page.get_text("text") or "").strip()
            if len(native) >= settings.ocr_min_chars:
                result.pages.append(PageText(page_no, native, ocr_used=False))
                continue
            # Sparse page -> likely scanned. Try OCR, keep the better of the two.
            ocr = _ocr_page(page, settings).strip()
            if len(ocr) > len(native):
                result.pages.append(PageText(page_no, ocr, ocr_used=True))
                if not ocr:
                    result.warnings.append(f"{filename} p.{page_no}: page illisible (OCR vide)")
            else:
                result.pages.append(PageText(page_no, native, ocr_used=False))
                if not native:
                    result.warnings.append(f"{filename} p.{page_no}: page illisible")
    return result


def extract_docx(content: bytes, filename: str) -> ExtractResult:
    import docx  # python-docx

    result = ExtractResult()
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"{filename}: DOCX illisible ({exc})")
        return result

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    # DOCX has no reliable page model; treat as a single logical page.
    result.pages.append(PageText(1, "\n".join(parts), ocr_used=False))
    return result


def extract_xlsx(content: bytes, filename: str) -> ExtractResult:
    import openpyxl

    result = ExtractResult()
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"{filename}: XLSX illisible ({exc})")
        return result

    for index, ws in enumerate(wb.worksheets, start=1):
        lines: list[str] = [f"[Feuille: {ws.title}]"]
        for row in ws.iter_rows(values_only=True):
            cells = [str(v) for v in row if v is not None]
            if cells:
                lines.append(" | ".join(cells))
        result.pages.append(PageText(index, "\n".join(lines), ocr_used=False))
    wb.close()
    return result


def extract_text(filename: str, content: bytes, settings: Settings) -> ExtractResult:
    """Dispatch on extension. Unknown types return an empty result + a warning."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf(content, filename, settings)
    if lower.endswith(".docx"):
        return extract_docx(content, filename)
    if lower.endswith(".xlsx"):
        return extract_xlsx(content, filename)
    return ExtractResult(warnings=[f"{filename}: format non supporté"])
